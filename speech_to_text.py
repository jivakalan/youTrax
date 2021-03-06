# -*- coding: utf-8 -*-
"""
Created on Fri Oct 23 20:02:41 2020

@author: kalan
"""

# pip install pytubeX
# pip install python-docx
# pip install moviepy

from pytube import YouTube
import os,sys,time,traceback,glob
from moviepy.editor import *
from datetime import datetime
from docx import Document
from google.oauth2 import service_account


def downloading_video(link):
        now = datetime.now() # current date and time
        timeNow = now.strftime("%m-%d-%Y at %H_%M_%S")

        yt = YouTube(link)
        time.sleep(5)
        video = yt.streams.filter(progressive=True, file_extension='mp4').order_by('resolution').desc().first()
        time.sleep(2)

        folderPath = os.path.join(os.getcwd(),timeNow)
        os.mkdir(folderPath)

        videoPath = os.path.join(os.getcwd(),timeNow,video.title)

        video.download(output_path = folderPath)

        return videoPath

def extracting_audio(videoPath,audioFormat):
    if os.path.exists(videoPath) == False:
        videoPath = glob.glob(os.path.join(os.path.dirname(videoPath),'*'))[0]

    audioPath = os.path.splitext(videoPath)[0]+'.%s'%audioFormat
    print('Video Path: ',videoPath)
    print('Audio Path: ',audioPath)

    video = VideoFileClip(videoPath)
    audio = video.audio
    audio.write_audiofile(audioPath)
    return audioPath





def create_bucket_class_location(bucket_name,credentials_path):
    """Create a new bucket in specific location with storage class"""
    # bucket_name = "your-new-bucket-name"
    from google.cloud import storage
  
    storage_client = storage.Client.from_service_account_json(credentials_path)
    
    # storage_client = storage.Client(credentials=credentials)
    
    try:
        bucket = storage_client.get_bucket(bucket_name)

        print("ID: {}".format(bucket.id))
    
    except:
        bucket = storage_client.bucket(bucket_name)
        bucket.storage_class = "COLDLINE"
        new_bucket = storage_client.create_bucket(bucket, location="us")

        print("Created bucket {} in {} with storage class {}".format(new_bucket.name, new_bucket.location, new_bucket.storage_class))

    return bucket_name

def uploading_file_to_gcs(bucket_name,audioPath,credentials_path):
    from google.cloud import storage

    storage_client = storage.Client.from_service_account_json(credentials_path)

    # storage_client = storage.Client(credentials=credentials)

    bucket = storage_client.get_bucket(bucket_name)

    directory = os.path.dirname(audioPath)

    destination = '%s/%s'%(os.path.basename(directory),os.path.basename(audioPath))
    
    blob = bucket.blob(destination)

    blob.upload_from_filename(audioPath)
    
    print("File {} uploaded to {}.".format( audioPath, destination))

    # URI of a file is formed like this
    # gs://<bucket_name>/<file_path_inside_bucket>

    uri = 'gs://%s/%s'%( bucket_name, destination)
    return uri

def tanscribe_from_audio_uri(videoPath, audioPath,model,punctuation,enable_word_time_offsets,sample_rate_hertz,credentials):
    """Transcribe the given audio file synchronously with
    the selected model."""
    from google.cloud import speech
    from google.cloud.speech import enums
    from google.cloud.speech import types

    client = speech.SpeechClient(credentials=credentials)

    # with open(audioPath, 'rb') as audio_file:
    #     content = audio_file.read()

    # audio = speech.types.RecognitionAudio(content=content)

    audio = types.RecognitionAudio(uri=audioPath)

    config = types.RecognitionConfig(
        # encoding=speech.enums.RecognitionConfig.AudioEncoding.EncodingUnspecified,
        encoding=speech.enums.RecognitionConfig.AudioEncoding.ENCODING_UNSPECIFIED,
        sample_rate_hertz=sample_rate_hertz,
        language_code='en-US',
        enable_word_time_offsets=enable_word_time_offsets,
        enable_automatic_punctuation=punctuation,
        model=model)

    operation = client.long_running_recognize(config, audio)

    print('Waiting for operation to complete...')
    result = operation.result(timeout=90)

    document = Document()

    for result in result.results:
        alternative = result.alternatives[0]
        print(u'Transcript: {}'.format(alternative.transcript))
        print('Confidence: {}'.format(alternative.confidence))

        document.add_paragraph(u'Transcript: {}'.format(alternative.transcript))
        document.add_paragraph('Confidence: {}'.format(alternative.confidence))

        if enable_word_time_offsets:
            start = False
            sentence = 'Sentence:\n' 
            for count, word_info in enumerate(alternative.words):
                word = word_info.word

                sentence+=word + ' '
                start_time = word_info.start_time
                end_time = word_info.end_time
                if start == False:
                    sentence_starting_time = '\nStart Time: {}'.format(start_time.seconds + start_time.nanos * 1e-9)
                    start = True
                if ('.' == word[-1]) or ('?' == word[-1]) or (count == (len(alternative.words)-1)):
                    sentence_ending_time = '\nEnd Time: {}'.format(end_time.seconds + end_time.nanos * 1e-9)

                    sentence_info = sentence + sentence_starting_time + sentence_ending_time
                    print(sentence_info) 

                    document.add_paragraph(sentence_info)
        
                    sentence = '\nSentence:\n' 
                    start = False

                

    document.save('%s.docx'%os.path.splitext(videoPath)[0])


def operation(link):
    #link = "https://www.youtube.com/watch?v=O16OBxFeKvc"

    videoPath = downloading_video(link)
    audioPath = extracting_audio(videoPath,'mp3')

    # audioPath = extracting_audio(videoPath,'wav')

    # Give path to the Speech to Text json file
    STT_credentials = service_account.Credentials.from_service_account_file('nth-mantra-233614-e06549201a47.json')
    # Give path to the Google Cloud Storage json file
    GCS_credentials_path = 'nth-mantra-233614-e06549201a47.json'

    bucket_name = 'bucket_jk'

    create_bucket_class_location(bucket_name,GCS_credentials_path)

    URI = uploading_file_to_gcs(bucket_name,audioPath,GCS_credentials_path)

    #Following is without punctuation and time offset
    tanscribe_from_audio_uri(videoPath, URI,'video',False,False,16000,STT_credentials)
    #Following is without time offset
    #tanscribe_from_audio_uri(videoPath, URI,'video',True,False,16000,STT_credentials)
    #Following is without punctuation 
    #tanscribe_from_audio_uri(videoPath, URI,'video',False,True,16000,STT_credentials)
    #Following is with both punctuation and time offset

    #tanscribe_from_audio_uri(videoPath, URI,'video',True,True,16000,STT_credentials)


if __name__ == "__main__":

    Link = input("https://www.youtube.com/watch?v=O16OBxFeKvc")
    try:
        operation(Link)
    except Exception as Error:
        traceback.print_exc()  
        print('Following Error has occured:\n',Error)
        
        Link='https://www.youtube.com/watch?v=O16OBxFeKvc'